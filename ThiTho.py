import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
import pdfplumber
import random
import re
import io
import time

# ================= GIAO DIỆN UI =================
st.set_page_config(page_title="ThiTho Pro X", layout="wide", page_icon="🎯")

st.markdown("""
<style>
.main .block-container { max-width: 95% !important; }
.question-box {
    background: #f8f9fa;
    padding: 18px;
    border-radius: 10px;
    border-left: 5px solid #0066cc;
    margin-bottom: 15px;
}
.question-text { font-size: 20px; font-weight: 700; color: #1e293b; }
.correct-highlight {
    background-color: #FFFF00; /* Bôi vàng hiển thị đáp án đúng */
    color: #FF0000; /* Chữ đỏ */
    font-weight: bold;
    padding: 2px 5px;
    border-radius: 4px;
}
</style>
""", unsafe_allow_html=True)

# ================= KHỞI TẠO STATE =================
for key in ["data_thi", "user_answers", "current_idx", "file_docx_clean", "next_trigger"]:
    if key not in st.session_state:
        st.session_state[key] = None if key in ["data_thi", "file_docx_clean"] else ({} if key == "user_answers" else (0 if key == "current_idx" else False))

# ================= HÀM CHUYÊN BIỆT LỌC DÒNG GIẢI THÍCH (NỀN XANH LÁ) =================
def check_is_explanation(line_text):
    """
    Nhận diện chính xác tất cả các kiểu dòng giải thích có nền xanh lá trong file HUBT để loại bỏ
    """
    txt = line_text.lower().strip()
    keywords = [
        "trong excel", "giải thích", "để di chuyển", "cấu trúc của", 
        "đáp án đúng", "hướng dẫn", "trong thiết kế", "công cụ số",
        "khi thiết kế", "địa chỉ tuyệt đối", "phần mềm access", "về côn",
        "địa chỉ ô", "mặc định", "hàm count", "hàm sum", "phần page", "phần report"
    ]
    return any(kw in txt for kw in keywords)

# ================= XỬ LÝ ĐỌC FILE WORD (.DOCX) SẠCH =================
def process_docx_clean(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    data = []
    current_q = None

    for para in doc.paragraphs:
        # Loại bỏ sạch chữ xám EduQuiz đứng dính hoặc rời
        text = re.sub(r'(?i)e\s*d\s*u\s*q\s*u\s*i\s*z', '', para.text).strip()
        if not text or text in ['Z', 'z', 'D', '--- PAGE ---']:
            continue

        # Phát hiện câu hỏi
        if text.lower().startswith("câu"):
            current_q = {"question": text, "options": [], "correct": None}
            data.append(current_q)
            continue

        # Bỏ hoàn toàn dòng giải thích (nền xanh)
        if check_is_explanation(text):
            continue

        # Cắt đáp án
        if current_q is not None:
            matches = re.findall(r'(\*?\s*[A-D]\.\s*.*?)(?=\s*\*?\s*[A-D]\.|$)', text)
            if matches:
                for m in matches:
                    m_clean = m.strip()
                    is_correct = m_clean.startswith("*")
                    opt_text = m_clean.replace("*", "").strip()
                    
                    if len(current_q["options"]) < 4:
                        current_q["options"].append(opt_text)
                        if is_correct: current_q["correct"] = opt_text

    return [q for q in data if len(q["options"]) >= 2]

# ================= XỬ LÝ ĐỌC FILE PDF (BÓC TÁCH LAYER CHỮ XÁM) =================
def process_pdf_clean(file_bytes):
    data = []
    
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        full_text = ""
        for page in pdf.pages:
            page_chars = page.chars
            if not page_chars: continue
            
            # Thuật toán Layer: Chỉ lấy chữ có kích thước chuẩn của câu hỏi/đáp án. 
            # Loại bỏ hoàn toàn các chữ xám khổng lồ (size > 15) ẩn dưới nền.
            lines_dict = {}
            for c in page_chars:
                if c["size"] > 15: 
                    continue # Bỏ qua chữ xám Watermark EduQuiz khổng lồ
                
                # Gom cụm các ký tự trên cùng một dòng dựa theo tọa độ dọc (top)
                top = round(c["top"], 1)
                found_line = False
                for t in lines_dict:
                    if abs(t - top) < 3:
                        lines_dict[t].append(c)
                        found_line = True
                        break
                if not found_line:
                    lines_dict[top] = [c]
            
            # Khôi phục văn bản sạch theo thứ tự từ trên xuống, từ trái sang
            for t in sorted(lines_dict.keys()):
                line_chars = sorted(lines_dict[t], key=lambda x: x["x0"])
                line_str = "".join([c["text"] for c in line_chars])
                full_text += line_str + "\n"

    # Dọn dẹp nốt tiêu đề lặp lại ở các đầu trang PDF
    full_text = re.sub(r'(?i)--- PAGE \d+ ---|TIN\s*3\s*-\s*HUBT\s*2026', '', full_text)
    full_text = re.sub(r'(?i)e\s*d\s*u\s*q\s*u\s*i\s*z', '', full_text)
    full_text = re.sub(r'Số trang:\s*\d+\s*Số câu hỏi:\s*\d+', '', full_text)
    full_text = re.sub(r'PHẦN\s*\d+:\s*[A-Z\s\(\)0-9\-]+', '', full_text)
    
    # Cắt văn bản thành các khối câu hỏi dựa vào "Câu [số]:" hoặc "Câu [số]."
    blocks = re.split(r'(?=Câu\s*\d+[:\.])', full_text)

    for block in blocks:
        lines = [x.strip() for x in block.split("\n") if x.strip()]
        if not lines: continue

        question = lines[0]
        if not question.lower().startswith("câu"): continue

        # Lọc bỏ hoàn toàn các dòng giải thích (chữ nền xanh)
        filtered_lines = []
        for line in lines[1:]:
            if check_is_explanation(line) or line.startswith("[Image") or line in ['Z', 'z', 'D']:
                continue
            filtered_lines.append(line)
            
        block_text = " ".join(filtered_lines)
        
        # Bóc tách 4 đáp án hệ A-B-C-D
        matches = re.findall(r'(\*?\s*[A-D]\.\s*.*?)(?=\s*\*?\s*[A-D]\.|$)', block_text)

        options = []
        correct = None

        for m in matches:
            m_clean = m.strip()
            is_correct = m_clean.startswith("*")
            clean_opt = m_clean.replace("*", "").strip()

            if len(options) < 4:
                options.append(clean_opt)
                if is_correct: correct = clean_opt

        if len(options) >= 2:
            data.append({"question": question, "options": options, "correct": correct})
            
    return data

# ================= HÀM XUẤT FILE WORD (.DOCX) SẠCH =================
def export_to_docx(data_list):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    for idx, item in enumerate(data_list, 1):
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(12)
        # Chuẩn hóa lại hiển thị số thứ tự câu từ 1 đến hết
        p_q.add_run(f"Câu {idx}: {item['question'].split(':', 1)[-1].strip() if ':' in item['question'] else item['question']}").bold = True

        for opt in item["options"]:
            p_o = doc.add_paragraph()
            p_o.paragraph_format.left_indent = Pt(20)
            if opt == item["correct"]:
                run_star = p_o.add_run("* ")
                run_star.font.color.rgb = RGBColor(255, 0, 0)
                run_star.bold = True
                run_txt = p_o.add_run(opt)
                run_txt.font.color.rgb = RGBColor(255, 0, 0)
                run_txt.bold = True
            else:
                p_o.add_run(opt)
                
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ================= THANH SIDEBAR CÀI ĐẶT =================
with st.sidebar:
    st.header("⚙️ THIẾT LẬP BỘ ĐỀ")
    uploaded_file = st.file_uploader("Tải lên đề thi (DOCX hoặc PDF)", type=["docx", "pdf"])
    shuffle_q = st.checkbox("Đảo thứ tự câu hỏi")
    shuffle_a = st.checkbox("Đảo thứ tự đáp án")

    if st.button("🚀 BẮT ĐẦU TRẮC NGHIỆM", use_container_width=True, type="primary"):
        if uploaded_file is not None:
            with st.spinner("Đang loại bỏ chữ xám và lọc bỏ phần giải thích nền xanh..."):
                file_bytes = uploaded_file.read()
                
                if uploaded_file.name.lower().endswith(".pdf"):
                    parsed_data = process_pdf_clean(file_bytes)
                else:
                    parsed_data = process_docx_clean(file_bytes)
                
                if parsed_data:
                    # Xuất file word sạch để lưu trữ trước khi xáo trộn đề
                    st.session_state.file_docx_clean = export_to_docx(parsed_data)
                    if shuffle_q: random.shuffle(parsed_data)
                    if shuffle_a:
                        for q in parsed_data: random.shuffle(q["options"])
                        
                    st.session_state.data_thi = parsed_data
                    st.session_state.user_answers = {}
                    st.session_state.current_idx = 0
                    st.session_state.next_trigger = False
                    st.rerun()
                else:
                    st.error("❌ Không thể bóc tách, kiểm tra lại định dạng file!")

    if st.session_state.data_thi:
        st.markdown("---")
        st.download_button(
            label="📥 TẢI FILE WORD SIÊU SẠCH (.DOCX)",
            data=st.session_state.file_docx_clean,
            file_name="De_Thi_Da_Loc_Sach.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        if st.button("🔄 Đổi bộ đề khác", use_container_width=True):
            st.session_state.data_thi = None
            st.session_state.file_docx_clean = None
            st.session_state.user_answers = {}
            st.session_state.current_idx = 0
            st.session_state.next_trigger = False
            st.rerun()

# ================= KHÔNG GIAN HIỂN THỊ CHÍNH =================
if st.session_state.data_thi:
    data = st.session_state.data_thi
    idx = st.session_state.current_idx
    item = data[idx]
    total = len(data)

    col1, col2, col3 = st.columns([1, 2.5, 1.2])

    # ===== CỘT TRÁI: THỐNG KÊ BIỂU ĐỒ =====
    with col1:
        st.write("### 📊 Tiến độ")
        done = len(st.session_state.user_answers)
        correct_count = sum(1 for i, ans in st.session_state.user_answers.items() if ans == data[i].get("correct"))
        st.write(f"Đã làm: {done}/{total}")
        st.write(f"Đúng: {correct_count} | Sai: {done - correct_count}")
        st.progress(done / total if total else 0)

    # ===== CỘT GIỮA: NỘI DUNG THI =====
    with col2:
        st.markdown(f"""
        <div class="question-box">
            <div class="question-text">Câu {idx+1}</div>
            <div>{item['question']}</div>
        </div>
        """, unsafe_allow_html=True)

        answered = idx in st.session_state.user_answers
        correct_ans = item.get("correct")
        selected_index = item["options"].index(st.session_state.user_answers[idx]) if answered else None

        choice = st.radio("Chọn đáp án:", item["options"], key=f"q_{idx}", index=selected_index, disabled=answered)

        if choice and not answered:
            st.session_state.user_answers[idx] = choice
            st.session_state.next_trigger = True
            st.rerun()

        if answered:
            if correct_ans == st.session_state.user_answers[idx]:
                st.success("ĐÚNG ✅")
            else:
                st.error("SAI ❌")
            st.markdown(f"**Đáp án đúng:** <span class='correct-highlight'>⭐ {correct_ans}</span>", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        if c1.button("⬅ Câu Trước") and idx > 0:
            st.session_state.current_idx -= 1
            st.rerun()
        if c2.button("Câu Tiếp ➡") and idx < total - 1:
            st.session_state.current_idx += 1
            st.rerun()

    # ===== CỘT PHẢI: MỤC LỤC ĐIỀU HƯỚNG =====
    with col3:
        st.write("### 📑 Mục lục")
        for i in range(0, total, 4):
            cols = st.columns(4)
            for j in range(4):
                k = i + j
                if k < total:
                    label = str(k + 1)
                    if k in st.session_state.user_answers:
                        label += " ✅" if st.session_state.user_answers[k] == data[k].get("correct") else " ❌"
                    btn_type = "primary" if k == idx else "secondary"
                    if cols[j].button(label, key=f"m_{k}", type=btn_type):
                        st.session_state.current_idx = k
                        st.rerun()

    # ===== TỰ ĐỘNG CHUYỂN CÂU (AUTO NEXT) =====
    if st.session_state.next_trigger:
        time.sleep(0.5)
        st.session_state.next_trigger = False
        if st.session_state.current_idx < total - 1:
            st.session_state.current_idx += 1
            st.rerun()
else:
    st.info("👈 Hãy tải file PDF/DOCX lên ở thanh bên trái để hệ thống tự động lọc sạch và chạy trắc nghiệm trực tuyến!")
